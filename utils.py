from docx import Document
from fpdf import FPDF
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import tempfile
import os
import unicodedata

def clean_text(text):
    return unicodedata.normalize('NFKD', text).encode('latin-1', 'ignore').decode('latin-1')


def generate_docx_v2(data):
    doc = Document()

    # Criar estilo de cabeçalho grande (usado no nome)
    styles = doc.styles
    if "TitleCustom" not in styles:
        title_style = styles.add_style("TitleCustom", WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = "Arial"
        title_style.font.size = Pt(20)
        title_style.font.bold = True

    # Nome
    doc.add_paragraph(data.get('full_name', ''), style="TitleCustom")

    # Contato
    contact_line1 = f"Cidade: {data.get('location', '')} | Celular: {data.get('phone', '')} | E-mail: {data.get('email', '')}"
    contact_line2 = f"LinkedIn: {data.get('linkedin', '')}"
    if data.get('github'):
        contact_line2 += f" | GitHub: {data['github']}"
    doc.add_paragraph(contact_line1)
    doc.add_paragraph(contact_line2)

    # Resumo Profissional
    doc.add_heading("Resumo Profissional", level=1)
    doc.add_paragraph(data.get('summary', ''))

    # Experiências
    doc.add_heading("Experiências", level=1)
    for exp in data.get('experiences', []):
        doc.add_paragraph(f"{exp.get('role', '')} - {exp.get('company', '')} ({exp.get('period', '')})", style='Heading 3')
        doc.add_paragraph(exp.get('description', ''))

    # Formação
    doc.add_heading("Formação", level=1)
    for edu in data.get('education', []):
        doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('institution', '')} ({edu.get('year', '')})")

    # Habilidades Técnicas
    doc.add_heading("Habilidades Técnicas", level=1)
    doc.add_paragraph(data.get('skills', ''))

    # Certificações
    certifications = data.get('certifications', '').strip()
    if certifications:
        doc.add_heading("Certificações", level=1)
        for cert in certifications.split("\n"):
            doc.add_paragraph(cert.strip())

    # Idiomas
    languages = data.get('languages', '').strip()
    if languages:
        doc.add_heading("Idiomas", level=1)
        for lang in languages.split("\n"):
            doc.add_paragraph(lang.strip())

    # Salvar e retornar caminho
    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_file.name)
    return tmp_file.name


def generate_docx(data):
    doc = Document()

    doc.add_heading(data.get('full_name', ''), 0)

    contact = f"Cidade: {data.get('location', '')} | Celular: {data.get('phone', '')} | E-mail: {data.get('email', '')} | LinkedIn: {data.get('linkedin', '')}"
    if data.get('github'):
        contact += f" | GitHub: {data['github']}"
    doc.add_paragraph(contact)

    doc.add_heading("Resumo Profissional", level=1)
    doc.add_paragraph(data.get('summary', ''))

    doc.add_heading("Experiência Profissional", level=1)
    for exp in data.get('experiences', []):
        doc.add_paragraph(f"{exp.get('role', '')} - {exp.get('company', '')} – {exp.get('period', '')}", style='ListBullet')
        doc.add_paragraph(exp.get('description', ''))

    doc.add_heading("Educação", level=1)
    for edu in data.get('education', []):
        doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('institution', '')} – {edu.get('year', '')}", style='ListBullet')

    doc.add_heading("Habilidades Técnicas", level=1)
    doc.add_paragraph(data.get('skills', ''))

    certifications = data.get('certifications', '').strip()
    if certifications:
        doc.add_heading("Certificações", level=1)
        for cert in certifications.split("\n"):
            doc.add_paragraph(cert.strip(), style='ListBullet')

    languages = data.get('languages', '').strip()
    if languages:
        doc.add_heading("Idiomas", level=1)
        for lang in languages.split("\n"):
            doc.add_paragraph(lang.strip(), style='ListBullet')

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_file.name)
    return tmp_file.name


def generate_pdf(data):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, data['full_name'], ln=True)

    pdf.set_font("Arial", '', 12)
    pdf.cell(0, 10, f"{data['location']} | {data['phone']} | {data['email']} | {data['linkedin']} | {data['github']}", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Professional Summary", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, data['summary'])

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Professional Experience", ln=True)
    pdf.set_font("Arial", '', 12)
    for exp in data['experiences']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, f"{exp['role']}", ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.cell(0, 10, f"{exp['company']} - {exp['period']}", ln=True)
        pdf.multi_cell(0, 10, exp['description'])

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Education", ln=True)
    pdf.set_font("Arial", '', 12)
    for edu in data['education']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, edu['degree'], ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.cell(0, 10, f"{edu['institution']} - {edu['year']}", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Technical Skills", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, data['skills'])

    if data['certifications']:
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, "Certifications", ln=True)
        pdf.set_font("Arial", '', 12)
        for cert in data['certifications']:
            pdf.cell(0, 10, cert, ln=True)

    if data['languages']:
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, "Languages", ln=True)
        pdf.set_font("Arial", '', 12)
        for lang in data['languages']:
            pdf.cell(0, 10, lang, ln=True)

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp_file.name)
    return tmp_file.name


def generate_docx_english(data):
    doc = Document()

    doc.add_heading(data.get('full_name', ''), 0)

    contact = f"Cidade: {data.get('location', '')} | Celular: {data.get('phone', '')} | E-mail: {data.get('email', '')} | LinkedIn: {data.get('linkedin', '')}"
    if data.get('github'):
        contact += f" | GitHub: {data['github']}"
    doc.add_paragraph(contact)

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(data.get('summary', ''))

    doc.add_heading("Professional Experience", level=1)
    for exp in data.get('experiences', []):
        doc.add_paragraph(f"{exp.get('role', '')}", style='ListBullet')
        doc.add_paragraph(f"{exp.get('company', '')} – {exp.get('period', '')}")
        doc.add_paragraph(exp.get('description', ''))

    doc.add_heading("Education", level=1)
    for edu in data.get('education', []):
        doc.add_paragraph(f"{edu.get('degree', '')}", style='ListBullet')
        doc.add_paragraph(f"{edu.get('institution', '')} – {edu.get('year', '')}")

    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph(data.get('skills', ''))

    certifications = data.get('certifications', '').strip()
    if certifications:
        doc.add_heading("Certifications", level=1)
        for cert in certifications.split("\n"):
            doc.add_paragraph(cert.strip(), style='ListBullet')

    languages = data.get('languages', '').strip()
    if languages:
        doc.add_heading("Languages", level=1)
        for lang in languages.split("\n"):
            doc.add_paragraph(lang.strip(), style='ListBullet')

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_file.name)
    return tmp_file.name


def generate_pdf_english(data):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, data['full_name'], ln=True)

    pdf.set_font("Arial", '', 12)
    pdf.cell(0, 10, f"{data['location']} | {data['phone']} | {data['email']} | {data['linkedin']} | {data['github']}", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Professional Summary", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, data['summary'])

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Professional Experience", ln=True)
    pdf.set_font("Arial", '', 12)
    for exp in data['experiences']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, f"{exp['role']}", ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.cell(0, 10, f"{exp['company']} - {exp['period']}", ln=True)
        pdf.multi_cell(0, 10, exp['description'])

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Education", ln=True)
    pdf.set_font("Arial", '', 12)
    for edu in data['education']:
        pdf.set_font("Arial", 'B', 12)
        pdf.cell(0, 10, edu['degree'], ln=True)
        pdf.set_font("Arial", '', 12)
        pdf.cell(0, 10, f"{edu['institution']} - {edu['year']}", ln=True)

    pdf.ln(5)
    pdf.set_font("Arial", 'B', 14)
    pdf.cell(0, 10, "Technical Skills", ln=True)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, data['skills'])

    if data['certifications']:
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, "Certifications", ln=True)
        pdf.set_font("Arial", '', 12)
        for cert in data['certifications']:
            pdf.cell(0, 10, cert, ln=True)

    if data['languages']:
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, "Languages", ln=True)
        pdf.set_font("Arial", '', 12)
        for lang in data['languages']:
            pdf.cell(0, 10, lang, ln=True)

    tmp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(tmp_file.name)
    return tmp_file.name
